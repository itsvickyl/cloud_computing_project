"""
extractor.py — Resume Text Extraction Module
=============================================

CLOUD / SERVERLESS DESIGN NOTES:
---------------------------------
This module is inherently serverless-friendly for the following reasons:

1. STATELESS, PURE-FUNCTION DESIGN:
   - Every function takes raw bytes in and returns plain text out.
   - No shared mutable state, no class instances, no database connections.
   - This makes it safe to invoke concurrently across multiple AWS Lambda
     instances or Google Cloud Function workers without race conditions.

2. IN-MEMORY I/O VIA BytesIO:
   - All file parsing uses BytesIO (in-memory byte streams), NOT the local
     filesystem. This is critical for serverless because:
       a) Lambda/Cloud Functions have limited /tmp storage (512 MB on Lambda).
       b) Files downloaded from S3 or GCS can be passed directly as bytes —
          no need to write them to disk first.
       c) Eliminates filesystem cleanup concerns between invocations.

3. LIGHTWEIGHT COLD-START FOOTPRINT:
   - pdfplumber and python-docx are small libraries compared to the ML models.
   - This module loads fast, contributing to sub-second cold starts even on
     Lambda's smallest memory tiers.

4. CLOUD EVENT INTEGRATION:
   - In a serverless pipeline, the Lambda handler (see lambda_handler.py or
     cloud_lambda_pseudo.py) downloads the resume from S3 using boto3, then
     passes the raw bytes directly to extract_text(). No local file paths needed.
   - Example flow:  S3 Upload Event → Lambda → boto3.get_object() → extract_text(bytes)

5. PAY-PER-USE SCALABILITY:
   - Text extraction is CPU-bound but fast (milliseconds per resume).
   - In a serverless model, you pay only for the compute time of each extraction.
   - 1,000 resumes processed = 1,000 short Lambda invocations, billed per 1ms.
"""

import pdfplumber
import docx
from io import BytesIO


def extract_text(file_bytes, file_type):
    """
    Route file bytes to the appropriate parser based on file type.

    SERVERLESS NOTE: This dispatcher pattern keeps the interface uniform —
    the Lambda handler doesn't need to know which parser to call, it just
    passes bytes + type. This decoupling is key for microservice architectures
    where the handler and extractor could even run as separate functions.
    """
    if file_type == "pdf":
        return extract_pdf(file_bytes)
    elif file_type == "docx":
        return extract_docx(file_bytes)
    else:
        raise ValueError("Unsupported file type. Please use 'pdf' or 'docx'.")


def extract_pdf(file_bytes):
    """
    Extract text from a PDF file provided as raw bytes.

    CLOUD NOTE: BytesIO wraps the in-memory bytes as a file-like object —
    this is exactly what you get after calling `s3_client.get_object()['Body'].read()`
    in AWS Lambda. No temp files, no disk I/O, fully ephemeral.

    In a Cloud Run or GCP Cloud Function deployment, the equivalent would be:
        blob = bucket.blob(key)
        file_bytes = blob.download_as_bytes()
        text = extract_pdf(file_bytes)
    """
    text = ""
    # pdfplumber accepts file-like objects — BytesIO provides exactly that
    # from raw bytes downloaded from cloud storage (S3, GCS, Azure Blob)
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def extract_docx(file_bytes):
    """
    Extract text from a DOCX file provided as raw bytes.

    CLOUD NOTE: python-docx's Document() constructor accepts file-like objects,
    making it equally compatible with in-memory bytes from cloud storage downloads.
    No filesystem access required — ideal for Lambda's read-only environment.
    """
    doc = docx.Document(BytesIO(file_bytes))
    return "\n".join([p.text for p in doc.paragraphs])
